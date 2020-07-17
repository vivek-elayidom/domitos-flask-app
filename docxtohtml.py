import docx
from docx.shared import Inches

import io
from docx import Document

from io import StringIO
import sys

#sys.setdefaultencoding('utf-8')

document = Document('one.docx')

all_paras = document.paragraphs
para_array = []
html_group = '*'
for para in all_paras:
    if(len(para.text)>0):
        #print((para.text).encode('utf-8'))
        print(para.style.name)
        para_array.append((para.text.encode('utf-8')).decode('utf-8'))
        textpiece = (para.text.encode('utf-8')).decode('utf-8')
        html = ''
        if(para.style.name == "Heading 7"):
            html = '<h4 class="mb-3">'+ textpiece + '</h4>'
        if(para.style.name == "List Paragraph"):
            html = '<li>'+ textpiece + '</li>'
        else:
            html = '<p>' + textpiece + '</p>'
        html_group = html_group + html

#print((html_group).encode('utf-8')).decode('utf-8')

text_file = open("output.html", "w")
n = text_file.write(html_group)
text_file.close()

'''for item in para_array:
    print(item)

print(len(para_array[12]))

c=0
for item in para_array:
    if(len(item) < 200):
        print("item -  " + item)
'''



#document = Document('docfiles/test.docx')


